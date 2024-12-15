import re
from PIL import Image
from pytesseract import image_to_string, pytesseract
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Укажите путь к исполняемому файлу Tesseract
pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

def extract_text(image_path, lang="rus"):
    """Извлечение текста с использованием Tesseract OCR."""
    config = "--oem 3 --psm 6 -l rus"  # PSM 6 для структурированного текста
    image = Image.open(image_path)
    return image_to_string(image, lang=lang, config=config)

def extract_table_from_text(text):
    """Извлечение таблицы из текста."""
    lines = text.splitlines()
    table_data = []
    start_table = False

    for line in lines:
        # Проверяем начало таблицы
        if re.search(r"№.*Температурные|Наименование", line, re.IGNORECASE):
            start_table = True
            continue

        if start_table:
            # Разделение строк таблицы на колонки
            row = [col.strip() for col in re.split(r"\s{2,}|\t", line) if col.strip()]
            if row:
                table_data.append(row)

    # Приведение строк к одной длине
    if table_data:
        max_cols = max(len(row) for row in table_data)
        normalized_data = [row + [""] * (max_cols - len(row)) for row in table_data]
        df = pd.DataFrame(normalized_data)

        # Автоматическое определение заголовков
        if re.match(r"^[А-Яа-яЁёA-Za-z\s]+$", df.iloc[0, 0]):
            df.columns = df.iloc[0]
            df = df[1:]

        return df

    return pd.DataFrame()

def save_table_to_docx(df, output_path):
    """Сохранение таблицы в Word-документ с форматированием."""
    doc = Document()
    table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])

    # Добавляем заголовки
    for col_idx, col_name in enumerate(df.columns):
        cell = table.cell(0, col_idx)
        cell.text = str(col_name)
        cell._element.get_or_add_tcPr().append(
            parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls("w")))
        )
        cell.paragraphs[0].runs[0].font.size = Pt(11)

    # Заполняем таблицу
    for row_idx, row in enumerate(df.itertuples(index=False), start=1):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            cell.text = str(value)
            cell.paragraphs[0].runs[0].font.size = Pt(10)

    # Устанавливаем границы таблицы
    for row in table.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = parse_xml(
                r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                r'<w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                r'<w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                r'<w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                r'</w:tcBorders>'
            )
            tcPr.append(tcBorders)

    doc.save(output_path)

def save_text_to_docx(text, output_path):
    """Сохранение текста в Word-документ."""
    doc = Document()
    doc.add_paragraph(text)
    doc.save(output_path)

def main():
    # Путь к изображению
    image_path = "JOSUKE.png"  # Замените на путь к вашему изображению
    output_doc_path = "output.docx"
    text_output_path = "output_text.docx"

    # Извлечение текста
    raw_text = extract_text(image_path, lang="rus")
    print(f"Распознанный текст:\n{raw_text}")

    # Извлечение таблицы
    df = extract_table_from_text(raw_text)

    if not df.empty:
        # Сохранение таблицы в Word
        save_table_to_docx(df, output_doc_path)
        print(f"Таблица сохранена в {output_doc_path}")

        # Отображение результата
        print("Распознанная таблица:")
        print(df)
    else:
        print("Таблица не найдена в тексте.")

    # Сохранение текста в Word
    save_text_to_docx(raw_text, text_output_path)
    print(f"Текст сохранен в {text_output_path}")

if __name__ == "__main__":
    main()
