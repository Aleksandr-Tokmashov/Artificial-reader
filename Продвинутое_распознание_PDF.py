'''
Этот код выполняет следующие действия:

Извлекает текст с помощью Tesseract.
Предобрабатывает изображения.
Конвертирует PDF в изображения.
Извлекает таблицы с помощью Tabula.
Удаляет строки с более чем 50% пропущенных значений и заменяет оставшиеся пропущенные значения на 0.
Добавляет таблицы в Word-документ.
Сохраняет результат в Word-файл.
'''
from pdf2image import convert_from_path
from pytesseract import image_to_string
from PIL import Image
from tabula import read_pdf
from docx import Document
import fitz  # PyMuPDF
import pandas as pd
import os

# Укажите путь к исполняемому файлу Poppler
poppler_path = r"C:\Program Files\poppler-24.08.0\Library\bin"

# Функция для извлечения текста с помощью Tesseract
def extract_text_with_tesseract(image, lang="rus"):
    try:
        return image_to_string(image, lang=lang)
    except Exception as e:
        print(f"Ошибка при распознавании текста: {e}")
        return ""

# Предобработка изображения
def preprocess_image(image):
    return image.convert("L")  # Конвертируем в оттенки серого для улучшения OCR

# Конвертация PDF в изображения
def convert_pdf_to_images(pdf_path):
    try:
        images = convert_from_path(pdf_path, poppler_path=poppler_path)
        print(f"Конвертировано {len(images)} изображений.")
        return images
    except Exception as e:
        print(f"Ошибка при конвертации PDF в изображения: {e}")
        return []

# Извлечение таблиц с помощью Tabula
def extract_tables_with_tabula(pdf_path, page):
    try:
        tables = read_pdf(pdf_path, pages=page, lattice=True, multiple_tables=True)
        return tables  # Возвращает список DataFrame
    except Exception as e:
        print(f"Ошибка при извлечении таблиц: {e}")
        return []

# Добавление таблиц в Word
def add_table_to_word(doc, df):
    rows, cols = df.shape
    table = doc.add_table(rows=rows + 1, cols=cols)

    # Заполнение заголовков таблицы
    for col_idx, col_name in enumerate(df.columns):
        cell = table.cell(0, col_idx)
        cell.text = str(col_name)
        cell.paragraphs[0].runs[0].font.bold = True  # Выделение жирным

    # Заполнение строк таблицы
    for row_idx, row in enumerate(df.itertuples(index=False), start=1):
        for col_idx, value in enumerate(row):
            table.cell(row_idx, col_idx).text = str(value)

# Функция для удаления строк с более чем 50% пропущенных значений и замены оставшихся на 0
def clean_table(df):
    # Удаление строк с более чем 50% пропущенных значений
    df = df.dropna(thresh=int(0.5 * len(df.columns)))

    # Замена оставшихся пропущенных значений на 0
    df = df.fillna(0)

    return df

# Основная функция обработки PDF и изображений
def process_file(file_path, output_doc_path, lang="rus", is_pdf=True):
    # Создаем новый документ Word
    doc = Document()

    if is_pdf:
        # Работаем с PDF
        # Получаем количество страниц
        total_pages = fitz.open(file_path).page_count

        # Конвертируем PDF в изображения для OCR
        images = convert_pdf_to_images(file_path)

        for page_num in range(1, total_pages + 1):
            print(f"Обработка страницы {page_num}...")

            # Добавляем заголовок для страницы
            doc.add_heading(f"Страница {page_num}", level=1)

            # Извлечение таблиц с помощью Tabula
            tables = extract_tables_with_tabula(file_path, page=page_num)

            # Если таблицы найдены, добавляем их в документ
            if tables:
                for table in tables:
                    cleaned_table = clean_table(table)
                    if not cleaned_table.empty:
                        doc.add_paragraph("Объединенная таблица:", style="Heading 2")
                        add_table_to_word(doc, cleaned_table)
            else:
                # Если таблиц не найдено, распознаем текст с помощью Tesseract
                if page_num - 1 < len(images):
                    image = preprocess_image(images[page_num - 1])
                    text = extract_text_with_tesseract(image, lang=lang)
                    doc.add_paragraph("Распознанный текст:")
                    doc.add_paragraph(text)
                else:
                    print(f"Изображение для страницы {page_num} не найдено.")

    else:
        # Работаем с изображением
        print("Обработка изображения...")
        image = Image.open(file_path)
        image = preprocess_image(image)
        text = extract_text_with_tesseract(image, lang=lang)
        doc.add_paragraph("Распознанный текст:")
        doc.add_paragraph(text)

    # Сохранение Word-файла
    doc.save(output_doc_path)
    print(f"Файл сохранен как {output_doc_path}")

# Запуск скрипта
if __name__ == "__main__":
    file_path = "тестик.pdf"  # Замените на путь к вашему PDF или изображению

    if not os.path.exists(file_path):
        print(f"Файл {file_path} не найден.")
    else:
        output_doc_path = "результат.docx"  # Замените на желаемый путь для Word-файла
        process_file(file_path, output_doc_path, lang="rus", is_pdf=True)  # Укажите is_pdf=False для изображений
